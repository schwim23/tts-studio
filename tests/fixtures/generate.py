"""
Generate sample input files (.txt, .md, .pdf, .docx) for tests and demos.

Used by the test suite (via conftest) and runnable standalone:
    python tests/fixtures/generate.py            # -> tests/fixtures/
    python tests/fixtures/generate.py /some/dir
"""

from __future__ import annotations

import sys
from pathlib import Path

# Known content the tests assert against. Includes:
#  - an acronym ("API") to verify cleaner expansion
#  - a "key" phrase to verify emphasis
#  - a question and multiple paragraphs to verify prosody/chunking
SAMPLE_TITLE = "The Future of Audio"
SAMPLE_PARAGRAPHS = [
    "This document explains a key idea about narration. "
    "A good API makes the whole pipeline simple.",
    "Is this the most important point? Absolutely. "
    "Never ignore the listener's experience.",
    "In conclusion, clear audio beats clever audio every time.",
]
SAMPLE_TEXT = "\n\n".join(SAMPLE_PARAGRAPHS)
SAMPLE_MD = f"# {SAMPLE_TITLE}\n\n" + "\n\n".join(
    f"**Note:** {p}" if i == 0 else p for i, p in enumerate(SAMPLE_PARAGRAPHS)
)


def _write_txt(path: Path) -> None:
    path.write_text(f"{SAMPLE_TITLE}\n\n{SAMPLE_TEXT}", encoding="utf-8")


def _write_md(path: Path) -> None:
    path.write_text(SAMPLE_MD, encoding="utf-8")


def _write_pdf(path: Path) -> None:
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    rect = pymupdf.Rect(72, 72, 500, 720)
    page.insert_textbox(rect, f"{SAMPLE_TITLE}\n\n{SAMPLE_TEXT}", fontsize=12)
    doc.save(str(path))
    doc.close()


def _write_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(SAMPLE_TITLE, level=1)
    for para in SAMPLE_PARAGRAPHS:
        doc.add_paragraph(para)
    doc.save(str(path))


def build_fixtures(dest: str | Path) -> dict[str, Path]:
    """Create all fixture files in `dest` and return {kind: path}."""
    dest = Path(dest)
    dest.mkdir(parents=True, exist_ok=True)
    paths = {
        "txt": dest / "sample.txt",
        "md": dest / "sample.md",
        "pdf": dest / "sample.pdf",
        "docx": dest / "sample.docx",
    }
    _write_txt(paths["txt"])
    _write_md(paths["md"])
    _write_pdf(paths["pdf"])
    _write_docx(paths["docx"])
    return paths


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else Path(__file__).parent
    created = build_fixtures(target)
    for kind, p in created.items():
        print(f"{kind:5} -> {p}  ({p.stat().st_size} bytes)")
